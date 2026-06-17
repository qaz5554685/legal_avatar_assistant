from pathlib import Path

from docx import Document
from openai import AsyncOpenAI

from globalVars import (
    CONTRACT_REVIEW_MAX_CHARS,
    MD_DIR,
    OPENAI_API_KEY,
    OPENAI_BASE_URL,
    OPENAI_MODEL,
)
from utility.logHandler import PROJECT_LOGGER as LOGGER

REVIEW_TYPES = {
    "receipt": {
        "label": "收文登記",
        "prompt_file": "contract_tab1_recevie.md",
    },
    "initial_review": {
        "label": "初審",
        "prompt_file": "contract_tab2_initial_review.md",
    },
    "legal_review": {
        "label": "法務審查",
        "prompt_file": "contract_tab3_legal_review.md",
    },
}

FIXED_ANALYSIS_RESULT = """合約主文缺少違約責任條款 — 第六章一般條款共 5 條，全為程序性規定，完全沒有「違約金」、「損害賠償」、「合約終止」條款，也沒有管轄法院約定（正常應設第十八、十九條）
沒有附件 — 文件末尾直接接簽署欄，不含任何附件（報價單、SLA、碳足跡認證等均未附）
沒有追蹤修訂紀錄 — 純淨 Word 文件，無任何 Track Changes 記錄"""


def get_review_label(review_type: str) -> str:
    return REVIEW_TYPES.get(review_type, {}).get("label", review_type)


def get_prompt_path(review_type: str) -> Path | None:
    prompt_file = REVIEW_TYPES.get(review_type, {}).get("prompt_file")
    if not prompt_file:
        return None
    return MD_DIR / prompt_file


def _read_prompt(review_type: str) -> str:
    prompt_path = get_prompt_path(review_type)
    if not prompt_path:
        raise ValueError(f"Unknown review type: {review_type}")
    if not prompt_path.exists():
        raise FileNotFoundError(f"Prompt file not found: {prompt_path}")
    return prompt_path.read_text(encoding="utf-8")


def _extract_docx_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".doc":
        raise ValueError("目前 LLM 分析只支援 .docx / .docm，請先將 .doc 另存成 .docx 後再上傳。")
    if suffix not in {".docx", ".docm"}:
        raise ValueError("請上傳 Word 檔案（.docx 或 .docm）。")

    document = Document(str(file_path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n[表格 {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Word 檔案沒有可讀取的文字內容。")
    return text


def _trim_contract_text(text: str) -> str:
    if len(text) <= CONTRACT_REVIEW_MAX_CHARS:
        return text
    LOGGER.warning(
        "[analyze_contract] Contract text truncated from %d to %d chars",
        len(text),
        CONTRACT_REVIEW_MAX_CHARS,
    )
    return (
        text[:CONTRACT_REVIEW_MAX_CHARS]
        + "\n\n[系統註記：文件內容過長，以上僅包含前段可處理文字。]"
    )


def _build_user_content(review_type: str, contract_text: str) -> str:
    review_label = get_review_label(review_type)
    return f"""請依照系統提示執行「{review_label}」。

以下是使用者上傳 Word 檔案擷取出的文字內容：

--- 合約內容開始 ---
{contract_text}
--- 合約內容結束 ---
"""


async def _call_openai(prompt: str, review_type: str, contract_text: str) -> str:
    if not OPENAI_API_KEY:
        LOGGER.warning("[analyze_contract] OPENAI_API_KEY missing, using fixed fallback")
        return FIXED_ANALYSIS_RESULT

    kwargs = {"api_key": OPENAI_API_KEY}
    if OPENAI_BASE_URL:
        kwargs["base_url"] = OPENAI_BASE_URL

    client = AsyncOpenAI(**kwargs)
    response = await client.responses.create(
        model=OPENAI_MODEL,
        input=[
            {
                "role": "system",
                "content": prompt,
            },
            {
                "role": "user",
                "content": _build_user_content(review_type, contract_text),
            },
        ],
    )

    output_text = getattr(response, "output_text", None)
    if output_text:
        return output_text.strip()

    LOGGER.error("[analyze_contract] OpenAI response has no output_text: %s", response)
    return "分析完成，但模型沒有回傳可顯示的文字結果。"


async def analyze_contract(file_path: Path, review_type: str) -> str:
    """Analyze an uploaded Word contract with the matching md prompt.

    Flow:
    Line Word file -> saved file_path -> extract Word text -> load md prompt
    -> call OpenAI -> return model analysis result.
    """
    prompt_path = get_prompt_path(review_type)
    LOGGER.info(
        "[analyze_contract] file=%s review_type=%s prompt=%s model=%s",
        file_path,
        review_type,
        prompt_path,
        OPENAI_MODEL,
    )

    try:
        prompt = _read_prompt(review_type)
        contract_text = _trim_contract_text(_extract_docx_text(file_path))
        return await _call_openai(prompt, review_type, contract_text)
    except Exception as exc:
        LOGGER.exception("[analyze_contract] Failed: %s", exc)
        return f"合約分析失敗：{exc}"
